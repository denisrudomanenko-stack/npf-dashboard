from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import List, Optional
from pydantic import BaseModel
import hashlib
import os
import shutil
from datetime import datetime
from pathlib import Path
import mimetypes

from app.database import get_db
from app.models.document import Document, DocumentStatus, DocumentType
from app.schemas.document import DocumentResponse
from app.services.document_service import document_service, CHUNK_CONFIG
from app.config import settings

router = APIRouter()

# File size limits
MAX_FILE_SIZE_MB = 30
MAX_INDEX_SIZE_MB = 10
MAX_FILE_SIZE = MAX_FILE_SIZE_MB * 1024 * 1024  # 30 MB
MAX_INDEX_SIZE = MAX_INDEX_SIZE_MB * 1024 * 1024  # 10 MB


def get_file_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


@router.get("/", response_model=List[DocumentResponse])
async def get_documents(
    skip: int = 0,
    limit: int = 100,
    status: DocumentStatus = None,
    document_type: DocumentType = None,
    db: AsyncSession = Depends(get_db)
):
    """Get all documents from database."""
    query = select(Document).where(Document.status != DocumentStatus.DELETED)
    if status:
        query = query.where(Document.status == status)
    if document_type:
        query = query.where(Document.document_type == document_type)
    query = query.order_by(Document.created_at.desc()).offset(skip).limit(limit)
    result = await db.execute(query)
    return result.scalars().all()


@router.get("/indexed")
async def get_indexed_documents():
    """Get all documents indexed in the RAG knowledge base."""
    return document_service.list_documents()


@router.get("/types")
async def get_document_types():
    """Get available document types with chunk configurations."""
    return {
        "types": list(CHUNK_CONFIG.keys()),
        "configs": CHUNK_CONFIG,
        "npf_types": ["regulation", "product", "presentation", "contract_template", "analytics", "faq"]
    }


@router.get("/stats")
async def get_document_stats():
    """Get knowledge base statistics."""
    return document_service.get_stats()


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(document_id: int, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(default="other"),
    title: str = Form(default=None),
    description: str = Form(default=None),
    auto_index: bool = Form(default=True),
    db: AsyncSession = Depends(get_db)
):
    """Upload and optionally index a document for RAG."""
    allowed_extensions = {'.pdf', '.docx', '.txt', '.md', '.xlsx', '.xls', '.csv'}
    file_ext = os.path.splitext(file.filename)[1].lower()

    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"File type {file_ext} not allowed. Allowed: {allowed_extensions}")

    contents = await file.read()

    # Check file size
    file_size = len(contents)
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large ({file_size / 1024 / 1024:.1f} MB). Maximum size: {MAX_FILE_SIZE_MB} MB"
        )

    # Disable auto-index for large files
    if file_size > MAX_INDEX_SIZE and auto_index:
        auto_index = False

    content_hash = get_file_hash(contents)

    # Check for duplicates by hash
    result = await db.execute(
        select(Document).where(Document.content_hash == content_hash)
    )
    existing = result.scalar_one_or_none()
    if existing:
        raise HTTPException(status_code=400, detail=f"Document already exists: {existing.title}")

    # Save file
    docs_dir = Path(settings.documents_directory)
    docs_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{content_hash[:16]}_{file.filename}"
    file_path = docs_dir / filename

    with open(file_path, "wb") as f:
        f.write(contents)

    # Create DB record
    doc_title = title or file.filename
    doc = Document(
        filename=filename,
        original_filename=file.filename,
        file_path=str(file_path),
        file_type=file_ext[1:],
        file_size=len(contents),
        document_type=DocumentType(document_type) if document_type in DocumentType.__members__ else DocumentType.OTHER,
        title=doc_title,
        description=description,
        content_hash=content_hash,
        status=DocumentStatus.ACTIVE
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    result = {
        "message": "Document uploaded",
        "document_id": doc.id,
        "filename": filename,
        "size": len(contents)
    }

    # Auto-index in RAG if requested
    if auto_index:
        try:
            index_result = await document_service.add_file(
                file_path=str(file_path),
                title=doc_title,
                doc_type=document_type
            )

            if "error" not in index_result:
                doc.indexed_at = datetime.utcnow()
                doc.chunk_count = index_result.get("chunks", 0)
                await db.commit()

                result["indexed"] = True
                result["chunks"] = index_result.get("chunks", 0)
                result["rag_doc_id"] = index_result.get("doc_id")
            else:
                result["indexed"] = False
                result["index_error"] = index_result.get("error")
        except Exception as e:
            result["indexed"] = False
            result["index_error"] = str(e)

    return result


@router.post("/add-text")
async def add_text_document(
    title: str = Form(...),
    content: str = Form(...),
    document_type: str = Form(default="other"),
    description: str = Form(default=None)
):
    """Add a text document directly to the knowledge base."""
    if not content.strip():
        raise HTTPException(status_code=400, detail="Content cannot be empty")

    result = await document_service.add_document(
        title=title,
        content=content,
        doc_type=document_type,
        metadata={"description": description} if description else None
    )

    return result


@router.post("/search")
async def search_documents(
    query: str = Form(...),
    top_k: int = Form(default=5),
    doc_type: str = Form(default=None)
):
    """Search the knowledge base."""
    if not query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    results = await document_service.search(
        query=query,
        top_k=top_k,
        doc_type=doc_type if doc_type else None
    )

    return {"query": query, "results": results, "count": len(results)}


class ReindexRequest(BaseModel):
    document_type: Optional[str] = None


@router.post("/{document_id}/reindex")
async def reindex_document(
    document_id: int,
    request: ReindexRequest = None,
    db: AsyncSession = Depends(get_db)
):
    """Re-index a document in the vector database with optional category change."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(status_code=400, detail="Document file not found")

    # Check file size for indexing
    file_size = os.path.getsize(doc.file_path)
    if file_size > MAX_INDEX_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large for indexing ({file_size / 1024 / 1024:.1f} MB). Maximum: {MAX_INDEX_SIZE_MB} MB"
        )

    # Update document type if provided
    new_doc_type = None
    if request and request.document_type:
        try:
            new_doc_type = DocumentType(request.document_type)
            doc.document_type = new_doc_type
        except ValueError:
            pass  # Keep existing type if invalid

    # Delete old chunks
    indexed_docs = document_service.list_documents()
    for indexed in indexed_docs:
        if doc.original_filename in indexed.get("title", "") or doc.title in indexed.get("title", ""):
            await document_service.delete_document(indexed["doc_id"])
            break

    # Re-index with new or existing type
    doc_type_value = (new_doc_type or doc.document_type).value if (new_doc_type or doc.document_type) else "other"
    index_result = await document_service.add_file(
        file_path=doc.file_path,
        title=doc.title,
        doc_type=doc_type_value
    )

    if "error" not in index_result:
        doc.indexed_at = datetime.utcnow()
        doc.chunk_count = index_result.get("chunks", 0)
        await db.commit()

    return {"message": "Document reindexed", "document_type": doc_type_value, **index_result}


@router.patch("/{document_id}/archive")
async def archive_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """Archive document and remove from search."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    doc.status = DocumentStatus.ARCHIVED
    await db.commit()

    return {"message": "Document archived"}


@router.delete("/{document_id}")
async def delete_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """Soft delete document and remove from RAG."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove from vector DB
    indexed_docs = document_service.list_documents()
    for indexed in indexed_docs:
        if doc.original_filename in indexed.get("title", "") or doc.title in indexed.get("title", ""):
            await document_service.delete_document(indexed["doc_id"])
            break

    doc.status = DocumentStatus.DELETED
    await db.commit()

    return {"message": "Document deleted"}


@router.delete("/indexed/{doc_id}")
async def delete_indexed_document(doc_id: str):
    """Delete a document from the RAG knowledge base by its RAG doc_id."""
    success = await document_service.delete_document(doc_id)
    if success:
        return {"message": f"Document {doc_id} deleted from knowledge base"}
    raise HTTPException(status_code=404, detail="Document not found in knowledge base")


@router.get("/{document_id}/download")
async def download_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """Download the original document file."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="Document file not found on disk")

    # Determine media type
    media_type, _ = mimetypes.guess_type(doc.file_path)
    if not media_type:
        media_type = "application/octet-stream"

    return FileResponse(
        path=doc.file_path,
        filename=doc.original_filename,
        media_type=media_type
    )


@router.post("/{document_id}/suggest-name")
async def suggest_document_name(document_id: int, db: AsyncSession = Depends(get_db)):
    """Analyze document and suggest a name based on content."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="Document file not found on disk")

    # Extract text from first pages
    text_content = ""
    try:
        if doc.file_type == "pdf":
            import fitz  # PyMuPDF
            pdf_doc = fitz.open(doc.file_path)
            max_pages = min(3, len(pdf_doc))
            for page_num in range(max_pages):
                page = pdf_doc[page_num]
                text_content += page.get_text() + "\n"
            pdf_doc.close()
        elif doc.file_type == "docx":
            from docx import Document as DocxDocument
            docx_doc = DocxDocument(doc.file_path)
            paragraphs = docx_doc.paragraphs[:30]  # First ~30 paragraphs
            text_content = "\n".join([p.text for p in paragraphs])
        elif doc.file_type in ["txt", "md"]:
            with open(doc.file_path, "r", encoding="utf-8") as f:
                text_content = f.read(5000)  # First 5000 chars
        else:
            text_content = doc.original_filename
    except Exception as e:
        text_content = doc.original_filename

    # Limit text for analysis
    text_content = text_content[:3000]

    # Use LLM to suggest name
    from app.services.ollama_service import ollama_service
    from app.config import settings

    prompt = f"""Проанализируй содержимое документа и предложи краткое, информативное название на русском языке.

Содержимое документа (первые страницы):
{text_content}

Требования к названию:
- Краткое (3-7 слов)
- Отражает суть документа
- На русском языке
- Без кавычек и специальных символов

Ответь ТОЛЬКО названием, без пояснений."""

    suggested_name = doc.original_filename  # Default fallback

    try:
        # Try Anthropic first if configured
        if settings.anthropic_api_key:
            import anthropic
            client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
            message = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=100,
                messages=[{"role": "user", "content": prompt}]
            )
            suggested_name = message.content[0].text.strip()
        elif await ollama_service.is_available():
            suggested_name = await ollama_service.generate(prompt)
            suggested_name = suggested_name.strip()
    except Exception as e:
        # Fallback: extract from filename
        suggested_name = doc.original_filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ")

    return {
        "document_id": doc.id,
        "current_title": doc.title,
        "original_filename": doc.original_filename,
        "suggested_name": suggested_name
    }


@router.patch("/{document_id}/rename")
async def rename_document(
    document_id: int,
    new_title: str = Form(...),
    db: AsyncSession = Depends(get_db)
):
    """Rename a document."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    old_title = doc.title
    doc.title = new_title
    await db.commit()

    return {
        "message": "Document renamed",
        "document_id": doc.id,
        "old_title": old_title,
        "new_title": new_title
    }


@router.patch("/{document_id}/category")
async def change_document_category(
    document_id: int,
    document_type: str = Form(...),
    db: AsyncSession = Depends(get_db)
):
    """Change document category."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    old_category = doc.document_type.value if doc.document_type else "other"

    try:
        new_doc_type = DocumentType(document_type)
    except ValueError:
        new_doc_type = DocumentType.OTHER

    doc.document_type = new_doc_type
    await db.commit()

    return {
        "message": "Category changed",
        "document_id": doc.id,
        "old_category": old_category,
        "new_category": new_doc_type.value
    }


@router.post("/{document_id}/ocr")
async def ocr_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """Perform OCR on a PDF document using Anthropic Claude Vision."""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.file_type != "pdf":
        raise HTTPException(status_code=400, detail="OCR only available for PDF files")

    if not doc.file_path or not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="Document file not found on disk")

    # Check if Anthropic API is configured
    if not settings.anthropic_api_key:
        raise HTTPException(
            status_code=503,
            detail="OCR requires Anthropic API key. Please configure ANTHROPIC_API_KEY in .env"
        )

    try:
        # Try to import and use pdf2image for PDF to image conversion
        import base64
        from pdf2image import convert_from_path
        import anthropic

        # Convert PDF pages to images
        images = convert_from_path(doc.file_path, dpi=150, first_page=1, last_page=5)

        if not images:
            raise HTTPException(status_code=400, detail="Could not convert PDF to images")

        # Use Anthropic Vision to extract text
        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

        extracted_texts = []
        for i, img in enumerate(images):
            # Convert PIL image to base64
            import io
            buffer = io.BytesIO()
            img.save(buffer, format="PNG")
            img_base64 = base64.b64encode(buffer.getvalue()).decode()

            message = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": img_base64
                                }
                            },
                            {
                                "type": "text",
                                "text": "Извлеки весь текст с этого изображения документа. Сохрани форматирование и структуру. Верни только текст без комментариев."
                            }
                        ]
                    }
                ]
            )

            extracted_texts.append(f"--- Страница {i+1} ---\n{message.content[0].text}")

        full_text = "\n\n".join(extracted_texts)

        # Save OCR result to a text file
        ocr_file_path = doc.file_path.replace(".pdf", "_ocr.txt")
        with open(ocr_file_path, "w", encoding="utf-8") as f:
            f.write(full_text)

        # Index the OCR text in RAG
        index_result = await document_service.add_document(
            title=f"{doc.title} (OCR)",
            content=full_text,
            doc_type=doc.document_type.value if doc.document_type else "other",
            metadata={"source": "ocr", "original_doc_id": doc.id}
        )

        return {
            "message": f"OCR выполнен успешно. Обработано страниц: {len(images)}",
            "pages_processed": len(images),
            "text_length": len(full_text),
            "ocr_file": ocr_file_path,
            "indexed": "error" not in index_result,
            "chunks": index_result.get("chunks", 0)
        }

    except ImportError as e:
        raise HTTPException(
            status_code=503,
            detail=f"OCR requires additional packages: pip install pdf2image anthropic. Error: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OCR failed: {str(e)}")
